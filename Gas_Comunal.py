import tkinter as tk
from tkinter import PhotoImage, ttk, messagebox 
from datetime import datetime
import sqlite3                                              # Necesario para abrir SQLITE3
import os 
import matplotlib.pyplot as plt
from docx import Document                                   # <-- Librería para generar archivos DOCX (Word)
import subprocess                                           # Módulo para abrir archivos automáticamente
import sys                                                   # Necesario para abrir archivos en diferentes OS y para el editor de código

#-------------------------------------------------------------------------------------------
#-------------------------------------------------------------------------------------------

# =======================================================
# ⚙️ CONFIGURACIÓN GLOBAL ⚙️ 
# =======================================================

# Ruta de la imagen (Ajustar si es necesario)
Ruta = r"Logo Gas.gif"
DB_NAME = "registro_gas.db" # Nombre del archivo de la base de datos

# Ruta de la carpeta (Usada ahora para guardar los DOCX)
RUTA_PDFS = r"Proyecto Gas\Achivos DOCX" 

imagen_global = None 
registro_id_counter = 1 
CODIGO_SECRETO = "0000" # <-- NUEVA CONSTANTE PARA LA CLAVE DE ACCESO

#-------------------------------------------------------------------------------------------
#-------------------------------------------------------------------------------------------

# =======================================================
#  FUNCIONES DE BASE DE DATOS SQLITE3
# =======================================================

def inicializar_db():
    """
    Conecta a la base de datos, crea la tabla 'clientes' si no existe 
    y verifica que el esquema esté actualizado (incluyendo 'cedula').
    """
    global registro_id_counter
    conn = None
    try:
        # 1. Conexión a la base de datos (crea el archivo si no existe)
        conn = sqlite3.connect(DB_NAME)
        cursor = conn.cursor()

        # 2. Crear la tabla si no existe (Schema actualizado con 'cedula')
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS clientes (
                id TEXT PRIMARY KEY,
                nombre TEXT NOT NULL,
                apellido TEXT NOT NULL,
                cedula TEXT NOT NULL,  -- CAMPO CÉDULA
                telefono TEXT NOT NULL,
                fecha_ingreso TEXT NOT NULL,
                cantidad_bombonas INTEGER NOT NULL
            )
        ''')
        conn.commit()
        
        # 3. REPARACIÓN DE ESQUEMA: Verificar y añadir la columna 'cedula' si falta.
        try:
            cursor.execute("SELECT cedula FROM clientes LIMIT 1")
        except sqlite3.OperationalError:
            print("Base de datos desactualizada. Añadiendo columna 'cedula'...")
            cursor.execute("ALTER TABLE clientes ADD COLUMN cedula TEXT NOT NULL DEFAULT ''")
            conn.commit()
            print("Columna 'cedula' añadida exitosamente.")

        # 4. Obtener el próximo ID consecutivo (buscando el ID más grande)
        cursor.execute("SELECT id FROM clientes ORDER BY id DESC LIMIT 1")
        ultimo_id = cursor.fetchone()

        if ultimo_id:
            ultimo_numero = int(ultimo_id[0].split('-')[1])
            registro_id_counter = ultimo_numero + 1
        else:
            registro_id_counter = 1 
            
    except sqlite3.Error as e:
        print(f"Error de base de datos al inicializar: {e}")
    finally:
        if conn:
            conn.close()

def insertar_registro(datos):
    """Inserta un nuevo registro en la tabla 'clientes'."""
    conn = None
    try:
        conn = sqlite3.connect(DB_NAME)
        cursor = conn.cursor()
        
        # Sentencia SQL con 7 parámetros (?, ?, ?, ?, ?, ?, ?)
        sql = '''
            INSERT INTO clientes (id, nombre, apellido, cedula, telefono, fecha_ingreso, cantidad_bombonas) 
            VALUES (?, ?, ?, ?, ?, ?, ?)
        '''
        cursor.execute(sql, datos)
        conn.commit()
        print("Registro guardado exitosamente en SQLite.")
        return True
    except sqlite3.Error as e:
        error_msg = f"Error al intentar guardar el registro:\n{e}"
        print(error_msg)
        messagebox.showerror("Error de Base de Datos", error_msg)
        return False
    finally:
        if conn:
            conn.close()

def obtener_todos_los_registros():
    """Recupera todos los registros de la tabla 'clientes'."""
    conn = None
    registros = []
    try:
        conn = sqlite3.connect(DB_NAME)
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM clientes")
        registros = cursor.fetchall()
    except sqlite3.Error as e:
        print(f"Error al obtener registros: {e}")
    finally:
        if conn:
            conn.close()
    return registros
    
def obtener_registros_por_mes():
    """
    Recupera los registros de clientes, agrupados y contados por Mes/Año.
    Retorna: (resumen_mensual, detalle_clientes)
    """
    conn = None
    resumen_mensual = []
    detalle_clientes = []
    try:
        conn = sqlite3.connect(DB_NAME)
        cursor = conn.cursor()
        
        # SQL para agrupar por mes y año. 
        # El formato de fecha es DD-MM-YYYY, así que el mes es SUBSTR(fecha_ingreso, 4, 2)
        # y el año es SUBSTR(fecha_ingreso, 7, 4).
        sql = '''
            SELECT 
                SUBSTR(fecha_ingreso, 4, 2) AS mes,  -- MM
                SUBSTR(fecha_ingreso, 7, 4) AS anio, -- YYYY
                COUNT(id) AS total_clientes,
                SUM(cantidad_bombonas) AS total_bombonas
            FROM clientes
            GROUP BY anio, mes
            ORDER BY anio DESC, mes DESC
        '''
        cursor.execute(sql)
        resumen_mensual = cursor.fetchall()
        
        # Obtener todos los registros completos para el detalle 
        cursor.execute("SELECT nombre, apellido, cedula, fecha_ingreso, cantidad_bombonas FROM clientes ORDER BY fecha_ingreso DESC")
        detalle_clientes = cursor.fetchall()

    except sqlite3.Error as e:
        print(f"Error al obtener resumen mensual de registros: {e}")
    finally:
        if conn:
            conn.close()
            
    return resumen_mensual, detalle_clientes

# FUNCIÓN DE DB PARA ESTADÍSTICAS GENERALES
def obtener_totales_generales():
    """Recupera el total de clientes y el total de bombonas registradas."""
    conn = None
    total_clientes = 0
    total_bombonas = 0
    try:
        conn = sqlite3.connect(DB_NAME)
        cursor = conn.cursor()
        
        # 1. Total Clientes (COUNT de IDs)
        cursor.execute("SELECT COUNT(id) FROM clientes")
        total_clientes = cursor.fetchone()[0]
        
        # 2. Total Bombonas (SUM de cantidad_bombonas)
        cursor.execute("SELECT SUM(cantidad_bombonas) FROM clientes")
        total_bombonas = cursor.fetchone()[0]
        
        # Si SUM es NULL (base de datos vacía), se convierte a 0
        if total_bombonas is None:
            total_bombonas = 0
            
    except sqlite3.Error as e:
        print(f"Error al obtener totales generales: {e}")
    finally:
        if conn:
            conn.close()
            
    return total_clientes, total_bombonas

def obtener_registro_por_cedula(cedula):
    """Recupera un registro por su cédula."""
    conn = None
    registro = None
    try:
        conn = sqlite3.connect(DB_NAME)
        cursor = conn.cursor()
        # Selecciona todos los campos necesarios
        cursor.execute("SELECT id, nombre, apellido, cedula, telefono, fecha_ingreso, cantidad_bombonas FROM clientes WHERE cedula=?", (cedula,))
        registro = cursor.fetchone()
    except sqlite3.Error as e:
        print(f"Error al buscar registro por cédula: {e}")
    finally:
        if conn:
            conn.close()
    return registro

def actualizar_registro(datos):
    """
    Actualiza un registro existente.
    El orden de los datos debe ser: (id, nombre, apellido, cedula, telefono, fecha, bombonas)
    """
    conn = None
    try:
        conn = sqlite3.connect(DB_NAME)
        cursor = conn.cursor()
        
        sql = '''
            UPDATE clientes 
            SET nombre=?, apellido=?, cedula=?, telefono=?, fecha_ingreso=?, cantidad_bombonas=?
            WHERE id=?
        '''
        # Reordenamos los datos para la sentencia SQL: (nombre, apellido, cedula, telefono, fecha, bombonas, id)
        # Índice:        (1,       2,        3,       4,       5,          6,        0)
        datos_sql = (datos[1], datos[2], datos[3], datos[4], datos[5], int(datos[6]), datos[0])
        
        cursor.execute(sql, datos_sql)
        conn.commit()
        print(f"Registro {datos[0]} modificado exitosamente en SQLite.")
        return True
    except sqlite3.Error as e:
        error_msg = f"Error al intentar modificar el registro:\n{e}"
        print(error_msg)
        messagebox.showerror("Error de Base de Datos", error_msg)
        return False
    finally:
        if conn:
            conn.close()

def eliminar_registro_db(id_registro):
    """Elimina un registro de la tabla 'clientes' por su ID."""
    conn = None
    try:
        conn = sqlite3.connect(DB_NAME)
        cursor = conn.cursor()
        cursor.execute("DELETE FROM clientes WHERE id=?", (id_registro,))
        conn.commit()
        print(f"Registro {id_registro} eliminado exitosamente de SQLite.")
        return True
    except sqlite3.Error as e:
        error_msg = f"Error al intentar eliminar el registro:\n{e}"
        print(error_msg)
        messagebox.showerror("Error de Base de Datos", error_msg)
        return False
    finally:
        if conn:
            conn.close()

# =======================================================
# FUNCIONES DE ACCIÓN Y UTILIDAD
# =======================================================

def validar_telefono(texto):
    """Valida que la entrada sea numérica y tenga un máximo de 11 dígitos."""
    if texto.isdigit() and len(texto) <= 11:
        return True
    if texto == "":
        return True
    return False
    
def nombre_mes(num_mes):
    """Convierte el número de mes (string 'MM') a su nombre en español."""
    meses = {
        '01': 'Enero', '02': 'Febrero', '03': 'Marzo', '04': 'Abril',
        '05': 'Mayo', '06': 'Junio', '07': 'Julio', '08': 'Agosto',
        '09': 'Septiembre', '10': 'Octubre', '11': 'Noviembre', '12': 'Diciembre'
    }
    return meses.get(num_mes, 'Mes Desconocido')

def verificar_clave(entry_clave, boton_codigo):
    """Verifica la clave y habilita o deshabilita el botón de acceso al código."""
    global CODIGO_SECRETO
    clave_ingresada = entry_clave.get()
    
    if clave_ingresada == CODIGO_SECRETO:
        boton_codigo.config(state='normal', text="Ver Código y Modificar [ACCESO CONCEDIDO]", bg="#32CD32") # Verde Lima
        messagebox.showinfo("Acceso", "Clave correcta. El botón 'Ver Código y Modificar' ha sido activado.")
    else:
        # Volver al estado inicial si se equivoca
        boton_codigo.config(state='disabled', text="Ver Código y Modificar", bg="#FFD700") 
        messagebox.showerror("Acceso Denegado", "Clave incorrecta. Acceso denegado.")
    
    # Limpiar el campo de entrada
    entry_clave.delete(0, tk.END) # Siempre limpiar por seguridad


# --------------------------------------------------------
# 🆕 FUNCIÓN CENTRAR VENTANA
# --------------------------------------------------------
def centrar_ventana(ventana, ancho, alto):
    """Calcula y aplica la posición para centrar la ventana en la pantalla."""
    # Asegura que la ventana se haya renderizado para obtener el tamaño de la pantalla
    ventana.update_idletasks() 
    
    screen_width = ventana.winfo_screenwidth()
    screen_height = ventana.winfo_screenheight()
    
    # Coordenadas X e Y para centrar
    x = (screen_width // 2) - (ancho // 2)
    y = (screen_height // 2) - (alto // 2)
    
    # Establece la geometría: ancho x alto + posición_x + posición_y
    ventana.geometry(f'{ancho}x{alto}+{x}+{y}')

# --------------------------------------------------------
# FUNCIÓN GENERAR REPORTE (DOCX)
# --------------------------------------------------------

def generar_reporte():
    """Genera un reporte DOCX con el resumen de registros por mes, el detalle y abre el archivo."""
    global RUTA_PDFS 
    
    # 1. Obtener los datos agrupados y el detalle
    resumen_mensual, detalle_clientes = obtener_registros_por_mes()

    if not resumen_mensual:
        messagebox.showinfo("Reporte", "No hay registros en la base de datos para generar el reporte.")
        return

    # 2. Configuración de python-docx
    document = Document() 
    
    # --- Estilos y Encabezados ---
    document.add_heading('REPORTE MENSUAL DE REGISTROS DE GAS', 0)
    document.add_paragraph(f"Fecha de Generación: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    document.add_paragraph('\n') # Salto de línea

    # 3. Resumen por Mes/Año
    document.add_heading('1. RESUMEN POR MES/AÑO', level=1)
    
    # Crear la tabla de resumen (4 columnas)
    tabla_resumen = document.add_table(rows=1, cols=4)
    tabla_resumen.style = 'Light Shading' 
    
    # Cabecera de la tabla de resumen
    header_cells = tabla_resumen.rows[0].cells
    header_cells[0].text = 'PERÍODO'
    header_cells[1].text = 'TOTAL CLIENTES'
    header_cells[2].text = 'TOTAL BOMBONAS'
    header_cells[3].text = 'TOTAL GENERAL' 
    
    # Poner la cabecera en negrita
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        
    total_general_clientes = 0
    total_general_bombonas = 0
    
    # Contenido de la tabla de resumen
    for mes_num, anio, clientes, bombonas in resumen_mensual:
        periodo = f"{nombre_mes(mes_num)} de {anio}"
        
        row_cells = tabla_resumen.add_row().cells
        row_cells[0].text = periodo
        row_cells[1].text = str(clientes)
        row_cells[2].text = str(bombonas)
        row_cells[3].text = str(clientes + bombonas) # Suma de clientes y bombonas para la fila (como total)
        
        total_general_clientes += clientes
        total_general_bombonas += bombonas

    # Fila de Totales
    total_row = tabla_resumen.add_row().cells
    total_row[0].text = 'TOTAL GENERAL'
    total_row[1].text = str(total_general_clientes)
    total_row[2].text = str(total_general_bombonas)
    total_row[3].text = str(total_general_clientes + total_general_bombonas)
    
    # Aplicar negrita a la fila de totales
    for cell in total_row:
        cell.paragraphs[0].runs[0].font.bold = True


    # 4. Detalle de Clientes
    document.add_paragraph('\n') # Salto de línea
    document.add_heading('2. DETALLE DE REGISTROS INDIVIDUALES', level=1)
    
    # Crear la tabla de detalle (5 columnas)
    tabla_detalle = document.add_table(rows=1, cols=5)
    tabla_detalle.style = 'Grid Table 4 Accent 1' 
    
    # Cabecera de la tabla de detalle
    detalle_header_cells = tabla_detalle.rows[0].cells
    detalle_header_cells[0].text = 'NOMBRE'
    detalle_header_cells[1].text = 'APELLIDO'
    detalle_header_cells[2].text = 'CÉDULA'
    detalle_header_cells[3].text = 'FECHA REG.'
    detalle_header_cells[4].text = 'BOMBONAS'
    
    # Poner la cabecera en negrita
    for cell in detalle_header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        
    # Contenido de la tabla de detalle
    for nombre, apellido, cedula, fecha_ingreso, cantidad_bombonas in detalle_clientes:
        row_cells = tabla_detalle.add_row().cells
        row_cells[0].text = nombre[:15] # Limitar a 15 caracteres
        row_cells[1].text = apellido[:15] # Limitar a 15 caracteres
        row_cells[2].text = cedula
        row_cells[3].text = fecha_ingreso
        row_cells[4].text = str(cantidad_bombonas)
        

    # 5. Guardar el DOCX y abrirlo
    nombre_base = f"Reporte_Mensual_Gas_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    ruta_completa = os.path.join(RUTA_PDFS, nombre_base)
    
    # Crear la carpeta si no existe
    try:
        if not os.path.exists(RUTA_PDFS):
            os.makedirs(RUTA_PDFS)
            print(f"Carpeta creada: {RUTA_PDFS}")
    except OSError as e:
        messagebox.showerror("Error de Carpeta", f"No se pudo crear la carpeta de destino: {RUTA_PDFS}\nError: {e}")
        return

    # Guardar el documento
    try:
        document.save(ruta_completa)
        print(f"Reporte DOCX guardado en: {ruta_completa}")
    except Exception as e:
        messagebox.showerror("Error de Guardado", f"Error al guardar el archivo DOCX: {e}")
        return

    # Abrir el archivo automáticamente (depende del OS)
    try:
        if sys.platform == "win32":
            os.startfile(ruta_completa)
        elif sys.platform == "darwin": # macOS
            subprocess.call(('open', ruta_completa))
        elif sys.platform.startswith('linux'): # Linux
            subprocess.call(('xdg-open', ruta_completa))
        else:
            messagebox.showinfo("Reporte Generado", f"Reporte generado exitosamente. No se pudo abrir automáticamente. Ubicación:\n{ruta_completa}")
            return # Termina si no se puede abrir

        messagebox.showinfo("Reporte Generado", "Reporte DOCX generado y abierto exitosamente.")
        
    except Exception as e:
         messagebox.showinfo("Reporte Generado", f"Reporte DOCX generado exitosamente, pero hubo un error al intentar abrirlo automáticamente:\n{e}")
         return


# --------------------------------------------------------
# FUNCIÓN PARA ABRIR VENTANA DE CONFIGURACIÓN
# --------------------------------------------------------

def abrir_ventana_configuracion(parent_window):
    """Crea la ventana de configuraciones del sistema."""
    
    # 1. Ocultar la ventana principal temporalmente
    parent_window.withdraw() 
    
    # 2. Configuración de la Ventana Toplevel
    ventana_config = tk.Toplevel(parent_window)
    # Transiencia (mantiene la ventana de Config sobre la principal)
    ventana_config.title("CONFIGURACIONES DEL SISTEMA")
    # Aplicar centrado
    centrar_ventana(ventana_config, 400, 550) 
    ventana_config.config(bg="#F0F8FF") # Azul claro

    # Protocolo de cierre para la 'X' de la barra de título (vuelve a mostrar la principal)
    ventana_config.protocol("WM_DELETE_WINDOW", lambda: [ventana_config.destroy(), parent_window.deiconify()])

    tk.Label(
        ventana_config,
        text="OPCIONES AVANZADAS",
        bg="#F0F8FF", 
        fg="#191970", 
        font=("Helvetica", 16, "bold")
    ).pack(pady=20)

    # --------------------------------------------------------
    # BOTONES DE CONFIGURACIÓN
    # --------------------------------------------------------

    # Botón Ver Manual
    tk.Button(
        ventana_config,
        text="Ver Manual del Sistema",
        command=lambda: mostrar_manual(parent_window), 
        bg="#ADD8E6", # Azul claro
        fg="black", 
        font=("Arial", 12, "bold"), 
        width=25
    ).pack(pady=20)
    
    # Botón Actualizar Config
    tk.Button(
        ventana_config,
        text="Actualizar Configuración",
        command=lambda: [ventana_config.withdraw(), actualizar_datos_sistema(parent_window)], # Oculta la ventana de Config y muestra la de Act.
        bg="#ADD8E6", # Azul claro
        fg="black", 
        font=("Arial", 12, "bold"), 
        width=25
    ).pack(pady=20)

    # Botón Gráficos
    tk.Button(
        ventana_config,
        text="Gráficos",
        command=lambda: [ventana_config.withdraw(), mostrar_graficos(parent_window)], # Oculta la ventana de Config y muestra la de Gráficos
        bg="#FFD700", fg="black", font=("Arial", 12, "bold"), width=25
    ).pack(pady=20)
    
    # ========================================================
    # 🔐 SEGURIDAD PARA ACCESO A CÓDIGO (NUEVO) 🔐
    # ========================================================
    
    # 1. Frame para el campo de clave y el botón 'Activar'
    seguridad_frame = tk.Frame(ventana_config, bg="#F0F8FF", padx=10, pady=5, relief=tk.GROOVE, borderwidth=2)
    seguridad_frame.pack(pady=(0, 10)) # Separación entre el botón de Gráficos y la seguridad
    
    # Etiqueta para el campo de clave
    tk.Label(
        seguridad_frame,
        text="Clave Única:",
        bg="#F0F8FF", 
        font=("Arial", 10)
    ).pack(side=tk.LEFT, padx=(0, 5))
    
    # Campo para la Clave Única (con show='*' para ocultar)
    clave_entry = tk.Entry(
        seguridad_frame,
        width=6, # Pequeño para la clave '0000'
        font=("Arial", 12),
        show="*" # Oculta los caracteres ingresados
    )
    clave_entry.pack(side=tk.LEFT, padx=5)
    
    # Botón "Activar" (pequeño)
    # Lo creamos aquí y asignaremos el comando al final
    activar_btn = tk.Button(
        seguridad_frame,
        text="Activar",
        # El comando se define después de crear codigo_btn para pasarlo como argumento
        bg="#00AA00", # Verde
        fg="white", 
        font=("Arial", 10, "bold"),
        width=6
    )
    activar_btn.pack(side=tk.LEFT, padx=5)
    
    # 2. Botón Ver Código y Modificar (INICIALMENTE DESHABILITADO)
    codigo_btn = tk.Button(
        ventana_config,
        text="Ver Código y Modificar",
        command=lambda: abrir_editor_codigo(parent_window), # No oculta ninguna ventana, solo Centra
        bg="#FFD700", fg="black", font=("Arial", 12, "bold"), width=25,
        state='disabled' # Deshabilitado por defecto
    )
    codigo_btn.pack(pady=20) 
    
    # 3. Asignar el comando al botón Activar después de crear codigo_btn
    activar_btn.config(command=lambda: verificar_clave(clave_entry, codigo_btn))
    
    # ========================================================
    # --------------------------------------------------------

    # Botón Cerrar
    tk.Button(
        ventana_config,
        text="Cerrar",
        # Comando para destruir la Toplevel y mostrar la ventana principal (Root)
        command=lambda: [ventana_config.destroy(), parent_window.deiconify()],
        bg="#FFCCCC", fg="black", font=("Arial", 12, "bold"), width=15
    ).pack(pady=20)

    return ventana_config
    
def actualizar_datos_sistema(parent_window):
    """Permite ver y modificar datos de configuración clave, como la ruta de la imagen."""
    global Ruta, RUTA_PDFS
    ventana_act = tk.Toplevel(parent_window)
    # Transiencia
    ventana_act.title("ACTUALIZACIÓN DE CONFIGURACIÓN")
    # Aplicar centrado
    centrar_ventana(ventana_act, 600, 300)
    ventana_act.config(bg="#FFFACD")
    
    tk.Label(
        ventana_act,
        text="Rutas de Archivos",
        bg="#FFFACD", 
        font=("Arial", 14, "bold")
    ).pack(pady=10)
    
    # Frame para los campos
    campos_frame = tk.Frame(ventana_act, bg="#FFFACD")
    campos_frame.pack(padx=20, pady=10)
    
    # Variables de control
    ruta_img_var = tk.StringVar(value=Ruta)
    ruta_docs_var = tk.StringVar(value=RUTA_PDFS)
    
    def guardar_config():
        """Guarda las nuevas rutas en las variables globales y las escribe en el archivo."""
        global Ruta, RUTA_PDFS
        
        # 1. Obtener y actualizar variables globales
        nueva_ruta_img = ruta_img_var.get()
        nueva_ruta_docs = ruta_docs_var.get()
        
        # Validaciones simples (ej. no vacíos)
        if not nueva_ruta_img or not nueva_ruta_docs:
            messagebox.showwarning("Error de Entrada", "Ambos campos de ruta son obligatorios.")
            return

        Ruta = nueva_ruta_img
        RUTA_PDFS = nueva_ruta_docs
        
        # 2. Intentar escribir los cambios en el archivo original (requiere conocimiento del script)
        # Esto es complejo y arriesgado. El editor de código ya maneja esto. 
        # Simplemente informamos que el cambio es temporal en la ejecución actual.
        
        messagebox.showinfo("Configuración", "Rutas de configuración actualizadas para esta sesión de ejecución. La imagen principal y la ruta de guardado DOCX cambiarán de inmediato.")

    # 1. Ruta de la Imagen
    tk.Label(campos_frame, text="Ruta Imagen (Logo Gas.gif):", bg="#FFFACD", font=("Arial", 10)).grid(row=0, column=0, sticky='w', padx=5, pady=5)
    tk.Entry(campos_frame, textvariable=ruta_img_var, width=50, font=("Arial", 10)).grid(row=0, column=1, padx=5, pady=5)

    # 2. Ruta de Guardado de Documentos
    tk.Label(campos_frame, text="Ruta Guardado DOCX (RELATIVA):", bg="#FFFACD", font=("Arial", 10)).grid(row=1, column=0, sticky='w', padx=5, pady=5)
    tk.Entry(campos_frame, textvariable=ruta_docs_var, width=50, font=("Arial", 10)).grid(row=1, column=1, padx=5, pady=5)
    
    # Botones de Acción
    botones_frame = tk.Frame(ventana_act, bg="#FFFACD")
    botones_frame.pack(pady=20)
    
    tk.Button(
        botones_frame,
        text="Guardar Configuración",
        command=guardar_config,
        bg="#4CAF50", fg="white", font=("Arial", 12, "bold"), width=20
    ).pack(side=tk.LEFT, padx=10)
    
    def cerrar_actualizacion():
        ventana_act.destroy()
        parent_window.deiconify() # Muestra la ventana principal

    tk.Button(
        botones_frame,
        text="Cerrar",
        command=cerrar_actualizacion,
        bg="#FFCCCC", fg="black", font=("Arial", 12, "bold"), width=15
    ).pack(side=tk.LEFT, padx=10)

    # Protocolo de cierre para la 'X' de la barra de título
    ventana_act.protocol("WM_DELETE_WINDOW", cerrar_actualizacion)
    
    return ventana_act

def abrir_editor_codigo(parent_window):
    """Abre un editor de texto con el código fuente del script para edición avanzada."""
    
    ruta_script = os.path.abspath(sys.argv[0]) # Obtiene la ruta del script actual
    
    # 1. Ocultar la ventana principal temporalmente
    parent_window.withdraw() 
    
    # 2. Configuración de la Ventana Toplevel
    ventana_codigo = tk.Toplevel(parent_window)
    ventana_codigo.title("EDITOR DE CÓDIGO FUENTE (Gas_Comunal.py)")
    centrar_ventana(ventana_codigo, 800, 600) 
    ventana_codigo.config(bg="#2E2E2E") # Fondo oscuro para código
    
    # --- Área de Edición ---
    codigo_frame = tk.Frame(ventana_codigo)
    codigo_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=(10, 5))
    
    # Scrollbar
    scrollbar_y = tk.Scrollbar(codigo_frame)
    scrollbar_y.pack(side=tk.RIGHT, fill=tk.Y)
    scrollbar_x = tk.Scrollbar(codigo_frame, orient=tk.HORIZONTAL)
    scrollbar_x.pack(side=tk.BOTTOM, fill=tk.X)
    
    # Text Widget para el código
    codigo_text = tk.Text(
        codigo_frame, 
        wrap=tk.NONE, # No wrapping
        yscrollcommand=scrollbar_y.set,
        xscrollcommand=scrollbar_x.set,
        undo=True, # Habilitar deshacer/rehacer
        bg="#1E1E1E", # Fondo más oscuro
        fg="#D4D4D4", # Letra clara
        insertbackground='white', # Cursor blanco
        font=("Consolas", 10)
    )
    codigo_text.pack(fill=tk.BOTH, expand=True)
    
    # Conectar Scrollbars
    scrollbar_y.config(command=codigo_text.yview)
    scrollbar_x.config(command=codigo_text.xview)
    
    # 3. Función Guardar Cambios
    def guardar_cambios():
        if not codigo_text.edit_modified():
            messagebox.showinfo("Guardar", "No se han realizado cambios para guardar.")
            return
            
        respuesta = messagebox.askyesno(
            "Confirmar Guardar",
            "⚠️ ¡ATENCIÓN! Guardar cambios puede introducir errores graves en el sistema.\n"
            "El programa debe REINICIARSE para aplicar los cambios.\n\n"
            "¿Está seguro de que desea SOBREESCRIBIR el archivo original con el código modificado?"
        )
        
        if respuesta:
            try:
                nuevo_contenido = codigo_text.get(1.0, tk.END)
                with open(ruta_script, 'w', encoding='utf-8') as f:
                    f.write(nuevo_contenido)
                    
                codigo_text.edit_modified(False)
                messagebox.showinfo("Éxito", "Código guardado exitosamente. Los cambios se aplicarán al **REINICIAR** el programa.")
                
            except Exception as e:
                messagebox.showerror("Error de Guardado", f"Error al guardar el archivo: {e}")

    # 4. Función de Cierre
    def cerrar_editor():
        if codigo_text.edit_modified():
            respuesta = messagebox.askyesno(
                "Cambios No Guardados",
                "Hay cambios sin guardar. ¿Desea salir sin guardar?"
            )
            if not respuesta:
                return # No cerrar
                
        ventana_codigo.destroy()
        parent_window.deiconify() # Muestra la ventana principal

    # --- Botones de Acción ---
    botones_inferiores_frame = tk.Frame(ventana_codigo, bg="#F0F8FF")
    botones_inferiores_frame.pack(fill=tk.X, pady=(5, 10), padx=10)
    
    tk.Button(
        botones_inferiores_frame,
        text="Guardar Cambios",
        command=guardar_cambios,
        bg="#FFA500", fg="black", font=("Arial", 12, "bold"), width=15
    ).pack(side=tk.LEFT, padx=10)
    
    tk.Button(
        botones_inferiores_frame,
        text="Cerrar Editor",
        command=cerrar_editor,
        bg="#FF6347", fg="white", font=("Arial", 12, "bold"), width=15
    ).pack(side=tk.RIGHT, padx=10)
    
    # Protocolo de cierre para la 'X' de la barra de título
    ventana_codigo.protocol("WM_DELETE_WINDOW", cerrar_editor)
    
    # --- Cargar el Contenido del Script ---
    try:
        with open(ruta_script, 'r', encoding='utf-8') as f:
            contenido = f.read()
        codigo_text.insert(tk.END, contenido)
        codigo_text.edit_modified(False) # Resetear la bandera de modificación inicial
    except Exception as e:
        codigo_text.insert(tk.END, f"ERROR al leer el archivo: {e}")
        messagebox.showerror("Error", f"Error al leer el archivo: {e}")
        return ventana_codigo

    return ventana_codigo

def mostrar_manual(parent_window):
    """Muestra una ventana con la información de manual del sistema."""
    parent_window.withdraw() # Oculta la ventana principal mientras se muestra el manual
    
    manual_texto = (
        "MANUAL DE OPERACIONES DEL SISTEMA DE REGISTRO DE GAS COMUNAL\n"
        "===========================================================\n\n"
        "▶️ REGISTRO NUEVO:\n"
        "   - Abre la ventana para ingresar la información de un nuevo cliente "
        "(Nombre, Apellido, Cédula, Teléfono, Cantidad de Bombonas).\n"
        "   - El sistema genera automáticamente la ID y la Fecha de Ingreso.\n"
        "   - La Cédula y la Cantidad de Bombonas deben ser números.\n\n"
        "▶️ VER REGISTROS:\n"
        "   - Muestra el listado completo de todos los clientes registrados en la Base de Datos.\n\n"
        "▶️ BUSCAR:\n"
        "   - Permite buscar un cliente por su Cédula.\n"
        "   - Una vez encontrado, se habilita la opción para Modificar o Eliminar el registro.\n\n"
        "▶️ CONFIGURACIÓN:\n"
        "   - Manual: Abre esta ventana de ayuda.\n"
        "   - Actualizar Configuración: Permite modificar rutas de configuración (ej. ruta de la imagen y ruta de guardado de DOCX).\n"
        "   - Gráficos: Muestra estadísticas clave (ej. total de clientes y bombonas).\n"
        "   - Ver Código: Abre el editor de código del propio sistema (Función Avanzada)."
    )
    messagebox.showinfo("Manual del Sistema", manual_texto)
    # Muestra la ventana de Config (parent_window) después de cerrar el messagebox
    parent_window.deiconify() 
    return None

def mostrar_graficos(parent_window):
    """Muestra estadísticas clave del sistema en una nueva ventana."""
    total_clientes, total_bombonas = obtener_totales_generales()
    
    ventana_stats = tk.Toplevel(parent_window)
    # Transiencia
    ventana_stats.title("ESTADÍSTICAS DEL SISTEMA")
    # Aplicar centrado
    centrar_ventana(ventana_stats, 400, 300) 
    ventana_stats.config(bg="#F0F8FF")

    tk.Label(
        ventana_stats,
        text="ESTADÍSTICAS GENERALES",
        bg="#F0F8FF", 
        fg="#191970", 
        font=("Helvetica", 16, "bold")
    ).pack(pady=20)

    tk.Label(
        ventana_stats,
        text=f"Total de Clientes Registrados: {total_clientes}",
        bg="#F0F8FF", 
        fg="black", 
        font=("Arial", 12)
    ).pack(pady=5)
    
    tk.Label(
        ventana_stats,
        text=f"Total de Bombonas Registradas: {total_bombonas}",
        bg="#F0F8FF", 
        fg="black", 
        font=("Arial", 12)
    ).pack(pady=5)
    
    # Botón Generar Reporte DOCX
    tk.Button(
        ventana_stats,
        text="Generar Reporte DOCX",
        command=generar_reporte,
        bg="#008080", # Verde azulado
        fg="white", 
        font=("Arial", 12, "bold"), 
        width=25
    ).pack(pady=20)
    
    def cerrar_stats():
        ventana_stats.destroy()
        parent_window.deiconify() # Muestra la ventana principal

    tk.Button(
        ventana_stats,
        text="Cerrar",
        command=cerrar_stats,
        bg="#FFCCCC", # Rojo claro
        fg="black", 
        font=("Arial", 12, "bold"), 
        width=15
    ).pack(pady=20)
    
    # Protocolo de cierre para la 'X' de la barra de título
    ventana_stats.protocol("WM_DELETE_WINDOW", cerrar_stats)
    
    return ventana_stats

# --------------------------------------------------------
# FUNCIÓN PARA ABRIR VENTANA DE REGISTRO
# --------------------------------------------------------

def abrir_ventana_registro(parent_window):
    """Abre la ventana para ingresar nuevos registros de clientes."""
    global registro_id_counter
    
    # 1. Ocultar la ventana principal temporalmente
    parent_window.withdraw() 

    # 2. Configuración de la Ventana Toplevel
    ventana_registro = tk.Toplevel(parent_window)
    ventana_registro.title("REGISTRAR NUEVO CLIENTE")
    centrar_ventana(ventana_registro, 500, 450)
    ventana_registro.config(bg="#F0F0F0")

    tk.Label(
        ventana_registro,
        text="INGRESO DE DATOS DEL CLIENTE",
        bg="#F0F0F0", 
        fg="#191970", 
        font=("Helvetica", 16, "bold")
    ).pack(pady=10)

    # Frame para los campos
    campos_frame = tk.Frame(ventana_registro, bg="#F0F0F0")
    campos_frame.pack(padx=20)

    # Variables de control
    id_var = tk.StringVar(value=f"REG-{registro_id_counter:04d}")
    nombre_var = tk.StringVar()
    apellido_var = tk.StringVar()
    cedula_var = tk.StringVar()
    telefono_var = tk.StringVar()
    fecha_var = tk.StringVar(value=datetime.now().strftime("%d-%m-%Y")) # Fecha actual
    bombonas_var = tk.StringVar()
    
    # Mapeo de campos para creación
    campos_info = [
        ("ID de Registro:", id_var, True),
        ("Nombre:", nombre_var, False),
        ("Apellido:", apellido_var, False),
        ("Cédula:", cedula_var, False),
        ("Teléfono:", telefono_var, False),
        ("Fecha de Ingreso:", fecha_var, True),
        ("Cantidad Bombonas:", bombonas_var, False)
    ]
    
    entries = {}
    
    # Validación del campo de teléfono
    vcmd_telefono = ventana_registro.register(validar_telefono)

    # Creación de campos (Label y Entry)
    for i, (label_text, var, readonly) in enumerate(campos_info):
        tk.Label(
            campos_frame, 
            text=label_text, 
            bg="#F0F0F0", 
            font=("Arial", 12)
        ).grid(row=i, column=0, padx=5, pady=5, sticky='w')
        
        entry = tk.Entry(
            campos_frame, 
            textvariable=var, 
            width=30, 
            font=("Arial", 12)
        )
        
        if label_text == "Teléfono:":
             entry.config(validate='key', validatecommand=(vcmd_telefono, '%P'))
        
        if readonly:
            entry.config(state='readonly', fg="gray")
            
        entry.grid(row=i, column=1, padx=5, pady=5, sticky='w')
        entries[label_text] = entry

    def guardar_registro():
        # 1. Recolectar datos
        id_reg = entries["ID de Registro:"].get()
        fecha_ingreso = entries["Fecha de Ingreso:"].get()
        
        datos = (
            id_reg,
            nombre_var.get().strip(),
            apellido_var.get().strip(),
            cedula_var.get().strip(),
            telefono_var.get().strip(),
            fecha_ingreso,
            bombonas_var.get().strip() # Dejar como string para validaciones
        )
        
        # 2. Validaciones
        if not all(datos[1:5]) or not datos[6]:
            messagebox.showwarning("Error de Entrada", "Todos los campos de Nombre, Apellido, Cédula, Teléfono y Bombonas son obligatorios.")
            return
            
        if not datos[3].isdigit():
            messagebox.showwarning("Error de Entrada", "El campo Cédula debe contener solo números.")
            return
            
        if not datos[6].isdigit() or int(datos[6]) <= 0:
            messagebox.showwarning("Error de Entrada", "La cantidad de bombonas debe ser un número entero mayor que cero.")
            return

        # 3. Inserción
        datos_para_db = (
            datos[0], datos[1], datos[2], datos[3], datos[4], datos[5], int(datos[6])
        )

        if insertar_registro(datos_para_db):
            messagebox.showinfo("Éxito", f"Cliente {datos[1]} {datos[2]} registrado exitosamente con ID: {datos[0]}")
            
            # Limpiar campos para el siguiente registro
            nombre_var.set("")
            apellido_var.set("")
            cedula_var.set("")
            telefono_var.set("")
            bombonas_var.set("")
            
            # Incrementar el contador global y actualizar la ID
            global registro_id_counter
            registro_id_counter += 1
            id_var.set(f"REG-{registro_id_counter:04d}")
        
    def cerrar_registro():
        ventana_registro.destroy()
        parent_window.deiconify() # Muestra la ventana principal

    # Frame para los botones
    botones_frame = tk.Frame(ventana_registro, bg="#F0F0F0")
    botones_frame.pack(pady=20)

    # Botón Guardar
    tk.Button(
        botones_frame,
        text="Guardar Registro",
        command=guardar_registro,
        bg="#4CAF50", # Verde
        fg="white", 
        font=("Arial", 12, "bold"), 
        width=15
    ).pack(side=tk.LEFT, padx=10)

    # Botón Cancelar/Cerrar
    tk.Button(
        botones_frame,
        text="Cerrar Ventana",
        command=cerrar_registro,
        bg="#FF6347", # Rojo
        fg="white", 
        font=("Arial", 12, "bold"), 
        width=15
    ).pack(side=tk.LEFT, padx=10)
    
    # Protocolo de cierre para la 'X' de la barra de título
    ventana_registro.protocol("WM_DELETE_WINDOW", cerrar_registro)
    
    return ventana_registro

# --------------------------------------------------------
# FUNCIÓN PARA ABRIR VENTANA DE LISTADO (Treeview)
# --------------------------------------------------------

def abrir_ventana_listado(parent_window):
    """Abre una ventana con un listado completo de los registros de clientes."""
    
    # 1. Ocultar la ventana principal temporalmente
    parent_window.withdraw() 
    
    # 2. Configuración de la Ventana Toplevel
    ventana_listado = tk.Toplevel(parent_window)
    ventana_listado.title("LISTADO COMPLETO DE CLIENTES")
    centrar_ventana(ventana_listado, 900, 500)
    ventana_listado.config(bg="#F0F0F0")

    tk.Label(
        ventana_listado,
        text="REGISTROS DE CLIENTES",
        bg="#F0F0F0", 
        fg="#191970", 
        font=("Helvetica", 16, "bold")
    ).pack(pady=10)

    # --- Configuración del Treeview ---
    tree_frame = tk.Frame(ventana_listado)
    tree_frame.pack(pady=10, padx=10, fill=tk.BOTH, expand=True)

    # Scrollbars
    scrollbar_y = tk.Scrollbar(tree_frame)
    scrollbar_y.pack(side=tk.RIGHT, fill=tk.Y)
    scrollbar_x = tk.Scrollbar(tree_frame, orient=tk.HORIZONTAL)
    scrollbar_x.pack(side=tk.BOTTOM, fill=tk.X)

    # Definir columnas
    columnas = ("id", "nombre", "apellido", "cedula", "telefono", "fecha_ingreso", "cantidad_bombonas")
    
    tree = ttk.Treeview(
        tree_frame, 
        columns=columnas, 
        show='headings', 
        yscrollcommand=scrollbar_y.set, 
        xscrollcommand=scrollbar_x.set
    )
    
    # Conectar Scrollbars al Treeview
    scrollbar_y.config(command=tree.yview)
    scrollbar_x.config(command=tree.xview)

    # Configurar las cabeceras de las columnas
    tree.heading("id", text="ID")
    tree.heading("nombre", text="Nombre")
    tree.heading("apellido", text="Apellido")
    tree.heading("cedula", text="Cédula")
    tree.heading("telefono", text="Teléfono")
    tree.heading("fecha_ingreso", text="Fecha Ingreso")
    tree.heading("cantidad_bombonas", text="Bombonas")

    # Configurar el ancho de las columnas
    tree.column("id", width=80, anchor=tk.CENTER)
    tree.column("nombre", width=120, anchor=tk.W)
    tree.column("apellido", width=120, anchor=tk.W)
    tree.column("cedula", width=100, anchor=tk.CENTER)
    tree.column("telefono", width=100, anchor=tk.CENTER)
    tree.column("fecha_ingreso", width=100, anchor=tk.CENTER)
    tree.column("cantidad_bombonas", width=80, anchor=tk.CENTER)

    tree.pack(fill=tk.BOTH, expand=True)

    # --- Cargar datos ---
    registros = obtener_todos_los_registros()
    
    # Limpiar cualquier dato previo
    for item in tree.get_children():
        tree.delete(item)
        
    for registro in registros:
        # Se insertan todos los campos. La cantidad de bombonas se asegura que sea string.
        tree.insert('', tk.END, values=(
            registro[0], # id
            registro[1], # nombre
            registro[2], # apellido
            registro[3], # cedula
            registro[4], # telefono
            registro[5], # fecha_ingreso
            str(registro[6]) # cantidad_bombonas
        ))
        
    # Función de cierre
    def cerrar_listado():
        ventana_listado.destroy()
        parent_window.deiconify() # Muestra la ventana principal

    # 1. Botón Cerrar
    tk.Button(
        ventana_listado,
        text="Cerrar",
        # Comando para destruir la Toplevel y mostrar la ventana principal (Root)
        command=cerrar_listado,
        bg="#FFCCCC", # Rojo claro
        fg="black", 
        font=("Arial", 12, "bold"), 
        width=15
    ).pack(pady=10)
    
    # Protocolo de cierre para la 'X' de la barra de título
    ventana_listado.protocol("WM_DELETE_WINDOW", cerrar_listado)
    
    return ventana_listado

# --------------------------------------------------------
# FUNCIÓN PARA ABRIR VENTANA DE BÚSQUEDA Y EDICIÓN
# --------------------------------------------------------

def abrir_ventana_busqueda(parent_window):
    """Permite buscar un registro por cédula para modificar o eliminar."""
    
    # 1. Ocultar la ventana principal temporalmente
    parent_window.withdraw() 

    # 2. Configuración de la Ventana Toplevel
    ventana_busqueda = tk.Toplevel(parent_window)
    ventana_busqueda.title("BÚSQUEDA Y MODIFICACIÓN")
    centrar_ventana(ventana_busqueda, 550, 650)
    ventana_busqueda.config(bg="#F0F8FF")

    tk.Label(
        ventana_busqueda,
        text="BUSCAR CLIENTE POR CÉDULA",
        bg="#F0F8FF", 
        fg="#191970", 
        font=("Helvetica", 16, "bold")
    ).pack(pady=10)
    
    # --- Frame de Búsqueda ---
    busqueda_frame = tk.Frame(ventana_busqueda, bg="#F0F8FF")
    busqueda_frame.pack(pady=10)
    
    tk.Label(
        busqueda_frame,
        text="Cédula a buscar:",
        bg="#F0F8FF", 
        font=("Arial", 12, "bold")
    ).pack(side=tk.LEFT, padx=5)
    
    cedula_busqueda_var = tk.StringVar()
    cedula_busqueda_entry = tk.Entry(
        busqueda_frame,
        textvariable=cedula_busqueda_var,
        width=20,
        font=("Arial", 12)
    )
    cedula_busqueda_entry.pack(side=tk.LEFT, padx=5)
    
    # Variable para almacenar el registro actual
    registro_actual = {} # Diccionario para (id, nombre, apellido, cedula, telefono, fecha, bombonas)

    # --- Variables de Control y Campos de Datos ---
    datos_frame = tk.Frame(ventana_busqueda, bg="#F0F8FF", pady=10, padx=20)
    datos_frame.pack(pady=10)
    
    # Variables de control
    id_var = tk.StringVar()
    nombre_var = tk.StringVar()
    apellido_var = tk.StringVar()
    cedula_var = tk.StringVar()
    telefono_var = tk.StringVar()
    fecha_var = tk.StringVar() 
    bombonas_var = tk.StringVar()
    
    # Mapeo de campos para creación
    campos_info = [
        ("ID:", id_var, True), 
        ("Nombre:", nombre_var, False),
        ("Apellido:", apellido_var, False),
        ("Cédula:", cedula_var, False),
        ("Teléfono:", telefono_var, False),
        ("Fecha Ingreso:", fecha_var, True), 
        ("Cantidad Bombonas:", bombonas_var, False)
    ]
    
    entries = {}
    
    # Validación del campo de teléfono
    vcmd_telefono = ventana_busqueda.register(validar_telefono)

    # Creación de campos (Label y Entry)
    for i, (label_text, var, readonly) in enumerate(campos_info):
        tk.Label(
            datos_frame, 
            text=label_text, 
            bg="#F0F8FF", 
            font=("Arial", 12)
        ).grid(row=i, column=0, padx=5, pady=5, sticky='w')
        
        entry = tk.Entry(
            datos_frame, 
            textvariable=var, 
            width=30, 
            font=("Arial", 10)
        )
        
        if label_text == "Teléfono:":
             entry.config(validate='key', validatecommand=(vcmd_telefono, '%P'))

        if readonly:
            entry.config(state='readonly')
            
        entry.grid(row=i, column=1, padx=5, pady=5, sticky='w')
        entries[label_text] = entry
        
    def limpiar_campos(habilitar_busqueda=True):
        id_var.set("")
        nombre_var.set("")
        apellido_var.set("")
        cedula_var.set("")
        telefono_var.set("")
        fecha_var.set("")
        bombonas_var.set("")
        registro_actual.clear()
        
        # Deshabilitar Edición (excepto ID y Fecha, que se cargan como readonly)
        entries["Nombre:"].config(state='disabled')
        entries["Apellido:"].config(state='disabled')
        entries["Cédula:"].config(state='disabled')
        entries["Teléfono:"].config(state='disabled')
        entries["Cantidad Bombonas:"].config(state='disabled')
        
        modificar_btn.config(state='disabled')
        eliminar_btn.config(state='disabled')
        
        if habilitar_busqueda:
            cedula_busqueda_entry.config(state='normal')
            cedula_busqueda_entry.delete(0, tk.END)

    def cargar_datos(registro):
        # Deshabilitar/habilitar para poder escribir con textvariable
        # Se hacen editable para la MODIFICACIÓN
        entries["Nombre:"].config(state='normal')
        entries["Apellido:"].config(state='normal')
        entries["Cédula:"].config(state='normal')
        entries["Teléfono:"].config(state='normal')
        entries["Cantidad Bombonas:"].config(state='normal')
        
        # Cargar variables
        id_var.set(registro[0])
        nombre_var.set(registro[1])
        apellido_var.set(registro[2])
        cedula_var.set(registro[3])
        telefono_var.set(registro[4])
        fecha_var.set(registro[5])
        bombonas_var.set(str(registro[6]))
        
        # Guardar el registro completo en la variable auxiliar
        registro_actual['datos'] = registro 
        
        # Habilitar botones de acción
        modificar_btn.config(state='normal')
        eliminar_btn.config(state='normal')
        
        # Deshabilitar el campo de búsqueda después de encontrar
        cedula_busqueda_entry.config(state='disabled')


    def buscar_registro():
        cedula = cedula_busqueda_var.get().strip()
        if not cedula or not cedula.isdigit():
            messagebox.showwarning("Entrada Inválida", "Por favor, ingrese una cédula válida (solo números).")
            return
            
        registro = obtener_registro_por_cedula(cedula)
        
        if registro:
            cargar_datos(registro)
            messagebox.showinfo("Éxito", f"Cliente {registro[1]} {registro[2]} encontrado.")
        else:
            limpiar_campos(habilitar_busqueda=True) # Limpia solo los campos de detalle y mantiene la búsqueda
            messagebox.showwarning("No Encontrado", f"No se encontró ningún cliente con la cédula: {cedula}")
        

    def modificar_registro():
        if not registro_actual:
            messagebox.showwarning("Advertencia", "No hay un registro cargado para modificar.")
            return

        # 1. Recolectar datos
        datos_nuevos = (
            id_var.get().strip(),
            nombre_var.get().strip(),
            apellido_var.get().strip(),
            cedula_var.get().strip(),
            telefono_var.get().strip(),
            fecha_var.get().strip(),
            bombonas_var.get().strip()
        )
        
        # 2. Validaciones (mismas que en el registro nuevo)
        if not all(datos_nuevos[1:5]) or not datos_nuevos[6]:
            messagebox.showwarning("Error de Entrada", "Todos los campos de Nombre, Apellido, Cédula, Teléfono y Bombonas son obligatorios.")
            return
            
        if not datos_nuevos[3].isdigit():
            messagebox.showwarning("Error de Entrada", "El campo Cédula debe contener solo números.")
            return
            
        if not datos_nuevos[6].isdigit() or int(datos_nuevos[6]) <= 0:
            messagebox.showwarning("Error de Entrada", "La cantidad de bombonas debe ser un número entero mayor que cero.")
            return

        # 3. Confirmación y Actualización
        respuesta = messagebox.askyesno(
            "Confirmar Modificación", 
            f"¿Está seguro de modificar el registro ID: {id_var.get()}?\n"
            f"Nuevo Nombre: {nombre_var.get()}\n"
            f"Nueva Cédula: {cedula_var.get()}"
        )
        
        if respuesta:
            if actualizar_registro(datos_nuevos):
                messagebox.showinfo("Éxito", f"Registro {id_var.get()} modificado correctamente.")
                limpiar_campos()
            # La función actualizar_registro ya maneja el error con messagebox

    def eliminar_registro():
        if not registro_actual:
            messagebox.showwarning("Advertencia", "No hay un registro cargado para eliminar.")
            return
            
        id_reg = id_var.get()
        nombre_reg = nombre_var.get()
        apellido_reg = apellido_var.get()

        respuesta = messagebox.askyesno(
            "Confirmar Eliminación", 
            f"⚠️ ¡ATENCIÓN! Esta acción es PERMANENTE.\n"
            f"¿Está seguro de eliminar el registro ID: {id_reg} del cliente {nombre_reg} {apellido_reg}?"
        )
        
        if respuesta:
            if eliminar_registro_db(id_reg):
                messagebox.showinfo("Éxito", f"Registro {id_reg} eliminado correctamente.")
                limpiar_campos()
            # La función eliminar_registro_db ya maneja el error con messagebox

    # Botón Buscar
    tk.Button(
        busqueda_frame,
        text="Buscar",
        command=buscar_registro,
        bg="#ADD8E6", fg="black", font=("Arial", 10, "bold")
    ).pack(side=tk.LEFT, padx=10)


    # --- Botones de Acción ---
    acciones_frame = tk.Frame(ventana_busqueda, bg="#F0F8FF")
    acciones_frame.pack(pady=20)

    modificar_btn = tk.Button(
        acciones_frame,
        text="Modificar",
        command=modificar_registro,
        bg="#FFD700", fg="black", font=("Arial", 12, "bold"), width=12,
        state='disabled'
    )
    modificar_btn.pack(side=tk.LEFT, padx=10)

    eliminar_btn = tk.Button(
        acciones_frame,
        text="Eliminar",
        command=eliminar_registro,
        bg="#FF6347", fg="white", font=("Arial", 12, "bold"), width=12,
        state='disabled'
    )
    eliminar_btn.pack(side=tk.LEFT, padx=10)
    
    # Botón Limpiar
    tk.Button(
        acciones_frame,
        text="Limpiar",
        command=limpiar_campos,
        bg="#98FB98", fg="black", font=("Arial", 12, "bold"), width=12
    ).pack(side=tk.LEFT, padx=10)

    # Función de cierre
    def cerrar_busqueda():
        ventana_busqueda.destroy()
        parent_window.deiconify()
        
    tk.Button(
        ventana_busqueda,
        text="Cerrar Ventana",
        command=cerrar_busqueda,
        bg="#FFCCCC", fg="black", font=("Arial", 12, "bold"), width=15
    ).pack(pady=10)

    # Inicializar campos deshabilitados al abrir la ventana
    limpiar_campos() 

    # Protocolo de cierre para la 'X' de la barra de título
    ventana_busqueda.protocol("WM_DELETE_WINDOW", cerrar_busqueda)
    
    return ventana_busqueda

# --------------------------------------------------------
# FUNCIÓN PRINCIPAL DE LA APLICACIÓN
# --------------------------------------------------------

def main():
    """Inicializa la base de datos y la ventana principal de la aplicación."""
    
    # Inicializar la base de datos y obtener el contador de ID
    inicializar_db() 
    
    # 1. Configuración de la Ventana Principal
    ventana_principal = tk.Tk()
    ventana_principal.title("Menu Principal")
    centrar_ventana(ventana_principal, 500, 550) 
    ventana_principal.config(bg="lightblue") 

    def abrir_ventana_con_ocultar(func_abrir):
        ventana_principal.withdraw()                # Oculta la ventana principal
        func_abrir(ventana_principal)               # Abre la nueva ventana (que debe tener el deiconify en su cierre)

    # --- Frame para el Título (Columna 0) ---
    titulo_frame = tk.Frame(ventana_principal, bg="lightblue")
    titulo_frame.grid(row=0, column=0, pady=10) 

    tk.Label(
        titulo_frame,
        text="SISTEMA DE GESTION DE GAS",
        bg="lightblue",
        fg="#191970", 
        font=("Helvetica", 18, "bold")
    ).pack()
    
    # ----------------------------------------------------------------------
    # --- Frame Contenedor de Imagen y Botones (Todo en columna 0, row 1) --
    # ----------------------------------------------------------------------
    
    contenedor_frame = tk.Frame(ventana_principal, bg="lightblue")
    contenedor_frame.grid(row=1, column=0, padx=10, pady=10) 
    
    # --- Sub-Frame para la Imagen 
    imagen_frame = tk.Frame(contenedor_frame, bg="lightblue", padx=10, pady=10)
    imagen_frame.pack(pady=10) 
    
    # Cargar y mostrar la imagen
    global imagen_global 
    try:
        ruta_imagen = os.path.join(os.path.dirname(os.path.abspath(__file__)), Ruta)
        imagen_cargada = PhotoImage(file=ruta_imagen)
        
        # 2. REDIMENSIONAR LA IMAGEN 
        nueva_anchura = 250                                 
        nueva_altura = 200 
        imagen_global = imagen_cargada.subsample(
            imagen_cargada.width() // nueva_anchura, 
            imagen_cargada.height() // nueva_altura
        )
        
        tk.Label(
            imagen_frame,
            image=imagen_global,
            bg="lightblue"
        ).pack()
    except Exception as e:
        tk.Label(imagen_frame, text="[Imagen no encontrada]", fg="red", bg="lightblue").pack()
        print(f"Error al cargar la imagen: {e}")

    # ------------------------------------------------------------
    # --- Sub-Frame para los Botones (Debajo de la Imagen) ---
    # ------------------------------------------------------------
    
    fila1_botones_frame = tk.Frame(contenedor_frame, bg="lightblue")
    fila1_botones_frame.pack(pady=5)

    fila2_botones_frame = tk.Frame(contenedor_frame, bg="lightblue")
    fila2_botones_frame.pack(pady=5)
    
    # --- Configuración de Fuente para Botones ---
    font_btn_grande = ("Arial", 14, "bold")

    # 1. Botón Registrar Nuevo (En la Fila 1, a la izquierda)
    registrar_btn = tk.Button(
        fila1_botones_frame,
        text="Registrar Nuevo",
        bg="#98FB98", 
        fg="black", 
        font=font_btn_grande, 
        width=15,
        command=lambda: abrir_ventana_con_ocultar(abrir_ventana_registro)
    )
    registrar_btn.pack(side=tk.LEFT, padx=10, pady=10)

    # 2. Botón Ver Registros 
    ver_registro_btn = tk.Button(
        fila1_botones_frame,
        text="Ver Registros",
        bg="#ADD8E6", 
        fg="black", 
        font=font_btn_grande,
        width=15, 
        command=lambda: abrir_ventana_con_ocultar(abrir_ventana_listado)
    )
    ver_registro_btn.pack(side=tk.LEFT, padx=10, pady=10)
    
    # 3. Botón Buscar 
    buscar_btn = tk.Button( 
        fila2_botones_frame,
        text="Buscar",
        bg="#ADD8E6",
        fg="black", 
        font=font_btn_grande, 
        width=15, 
        command=lambda: abrir_ventana_con_ocultar(abrir_ventana_busqueda)
    )
    buscar_btn.pack(side=tk.LEFT, padx=10, pady=10)
    
    # 4. Botón Configuracion 
    config_btn = tk.Button(
        fila2_botones_frame,
        text="Configuracion",
        bg="#ADD8E6", 
        fg="black", 
        font=("Arial", 12, "bold"), 
        width=15, # Reducido
        command=lambda: abrir_ventana_con_ocultar(abrir_ventana_configuracion) 
    )
    config_btn.pack(side=tk.LEFT, padx=10, pady=10) 
    
    # ------------------------------------------------------------
    # --- Frame para el Botón Salir (Centrado y Pequeño) ---
    # ------------------------------------------------------------
    def cerrar_aplicacion():
        ventana_principal.destroy()
        
    salir_frame = tk.Frame(ventana_principal, bg="lightblue")
    # Colocado en la última fila (row=2) de la ventana principal
    salir_frame.grid(row=2, column=0, pady=10)

    # 5. Botón Salir (Último y centrado)
    salir_btn = tk.Button(
        salir_frame,
        text="Salir",
        bg="#FF6347", 
        fg="white", 
        font=("Arial", 11, "bold"), # Fuente un poco más pequeña
        width=10, # Más pequeño
        command=cerrar_aplicacion
    )
    salir_btn.pack() # Pack lo centra dentro del frame

    # Ejecutar el bucle principal de tkinter
    ventana_principal.mainloop()

if __name__ == "__main__":
    main()